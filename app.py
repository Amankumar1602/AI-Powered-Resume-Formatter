
import streamlit as st
from openai import OpenAI
import pdfplumber
import docx
import os
from io import BytesIO
from docx import Document

# Fetch API key securely from Streamlit Secrets
OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
client = OpenAI(api_key=OPENAI_API_KEY)

# Function to extract text from a PDF
def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text.strip()

# Function to extract text from a DOCX file
def extract_text_from_docx(uploaded_file):
    doc = docx.Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])


def rewrite_resume(text):
    prompt = f"""
    Given the following resume text, extract the key details and rewrite it in the following structured format Don't add any extra fields only follow the format:

    Name of the person who's resume is uploaded in bold

    **Professional Summary:**  
    - - A concise overview of the candidate's experience and skills

    -------------------------------------------- add this to seperate Professional Summary and Primary Skillsets field

    **Primary Skillsets:**  
    - A bullet-point list of the candidate's main technical and soft skills.

    -------------------------------------------- add this to seperate Primary Skillsets and Education field

    **Education:**  
    - Degree, institution, and year of graduation.

    -------------------------------------------- add this to seperate Eduction and projects field

    **Projects:**  
    - List of key projects. 
      with each projects in the following format:
      Name of the project in bold
      - Role: The position or role of the candidate during the project
      - Enviroment: should contain the tech and editor used for the project
      - Description: about the project
      - Responsibilities: the work done by the candidate 

      ------------------------------------------------- add this to seperate each project 

    Here is the original resume text:
    {text}

    Please ensure clarity, conciseness, and professional formatting.
    """

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=3000,
        temperature=0.5
    )
    return response.choices[0].message.content.strip()

# Function to generate DOCX file
def generate_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    
    st.title("📄 AI-Powered Resume Formatter")

    uploaded_file = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])

    if uploaded_file is not None:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()

        with st.spinner("Processing..."):
            if file_extension == ".pdf":
                resume_text = extract_text_from_pdf(uploaded_file)
            elif file_extension == ".docx":
                resume_text = extract_text_from_docx(uploaded_file)
            else:
                st.error("Unsupported file format. Please upload a PDF or DOCX file.")
                st.stop()

            # Format the resume using LLM
            formatted_resume = rewrite_resume(resume_text)

        st.subheader("📌 Formatted Resume")
        st.markdown(formatted_resume)

        # Option to download formatted resume as a DOCX file
        docx_buffer = generate_docx(formatted_resume)
        st.download_button(
            label="📥 Download Formatted Resume (DOCX)",
            data=docx_buffer,
            file_name="formatted_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
